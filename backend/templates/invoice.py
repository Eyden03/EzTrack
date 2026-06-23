import os
from docx import Document
from datetime import datetime

TEMPLATES_DIR = os.path.dirname(__file__)

_TEMPLATES = {
    "sigla": "EzTrack_Non-VAT_Invoice_Template_Simula_Sigla.docx",
    "unlad": "EzTrack_VAT_Invoice_Template_Unlad.docx",
}

def _find_and_replace(doc, placeholder, value):
    value = str(value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            return
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return

def _find_table_by_header(doc, header_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if header_text.lower() in cell.text.lower():
                    return table
    return None

def _fill_items_table(table, items):
    if not table:
        return
    header_rows = 1
    data_rows = len(table.rows) - header_rows
    for idx, item in enumerate(items):
        if idx < data_rows:
            row = table.rows[header_rows + idx]
        else:
            row = table.add_row()
        amt = item["qty"] * item["unit_price"]
        cells = row.cells
        if len(cells) >= 4:
            cells[0].text = item.get("description", "")
            cells[1].text = str(item["qty"])
            cells[2].text = f"₱{item['unit_price']:,.2f}"
            cells[3].text = f"₱{amt:,.2f}"
        elif len(cells) >= 2:
            cells[0].text = item.get("description", "")
            cells[1].text = f"₱{amt:,.2f}"

def _amount_to_words(amount):
    if amount == 0:
        return "Zero Pesos Only"
    ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
            "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
            "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    def _convert(n):
        if n < 20:
            return ones[n]
        if n < 100:
            return tens[n // 10] + (" " + ones[n % 10] if n % 10 else "")
        if n < 1000:
            return ones[n // 100] + " Hundred" + (" " + _convert(n % 100) if n % 100 else "")
        if n < 1000000:
            return _convert(n // 1000) + " Thousand" + (" " + _convert(n % 1000) if n % 1000 else "")
        return ""
    whole = int(amount)
    cents = round((amount - whole) * 100)
    result = _convert(whole) + " Pesos"
    if cents:
        result += " and " + _convert(cents) + " Centavos"
    return result + " Only"

def build(items, customer, payment, ctx, tier, discount=0, withholding_tax=0, notes=""):
    tmpl = _TEMPLATES.get(tier, _TEMPLATES["sigla"])
    doc = Document(os.path.join(TEMPLATES_DIR, tmpl))
    biz = ctx.get("business") or {}
    biz_name = biz.get("name", "My Business")
    biz_city = biz.get("city", "")
    contact = biz.get("contact", "09XX-XXX-XXXX")

    _find_and_replace(doc, "[Your Business Name]", biz_name)
    _find_and_replace(doc, "[Business Address", biz_city)
    _find_and_replace(doc, "09XX-XXX-XXXX", contact)
    _find_and_replace(doc, "[Month DD, YYYY]", datetime.now().strftime("%B %d, %Y"))
    _find_and_replace(doc, "0000001", str(ctx.get("nextDocNumber", 1)))
    _find_and_replace(doc, "[Customer Name]", customer.get("name", ""))

    items_table = _find_table_by_header(doc, "ITEM")
    if not items_table:
        items_table = _find_table_by_header(doc, "QTY")
    _fill_items_table(items_table, items)

    subtotal = sum(i["qty"] * i["unit_price"] for i in items)
    total_due = subtotal - discount - withholding_tax

    _find_and_replace(doc, "₱[0.00]", f"₱{total_due:,.2f}")
    _find_and_replace(doc, "[0.00]", f"{subtotal:,.2f}")
    _find_and_replace(doc, "₱[0.00]", f"₱{discount:,.2f}")
    _find_and_replace(doc, "₱[0.00]", f"₱{withholding_tax:,.2f}")

    _find_and_replace(doc, "Total Sales:", f"Total Sales: ₱{subtotal:,.2f}")
    _find_and_replace(doc, "TOTAL AMOUNT DUE", f"TOTAL AMOUNT DUE")
    _find_and_replace(doc, "₱[0.00]", f"₱{total_due:,.2f}")

    if payment.get("method"):
        _find_and_replace(doc, "Cash", payment["method"])
    if payment.get("status"):
        _find_and_replace(doc, "Paid in Full", "✓ " + payment["status"])
        _find_and_replace(doc, "Partial", "")
        _find_and_replace(doc, "☐", "✓")

    if tier == "unlad":
        buyer_tin = customer.get("tin", "000-000-000-000")
        buyer_addr = customer.get("address", "")
        _find_and_replace(doc, "[Buyer's Registered Name]", customer.get("name", ""))
        _find_and_replace(doc, "[Buyer's TIN", buyer_tin)
        _find_and_replace(doc, "[Buyer's Address]", buyer_addr)
        vat_sales = subtotal / 1.12 if subtotal else 0
        vat_amt = subtotal - vat_sales
        _find_and_replace(doc, "₱[0.00]", f"₱{vat_sales:,.2f}")
        _find_and_replace(doc, "₱[0.00]", f"₱{vat_amt:,.2f}")
        _find_and_replace(doc, "[e.g. Please settle", notes or "Please settle payment within 7 days of invoice date.")

    if notes:
        _find_and_replace(doc, "[e.g. Please settle", notes)

    return doc
